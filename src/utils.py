# # src/utils.py
# import os, re, sys, types
# from pathlib import Path
# import pdfplumber
# from docx import Document as DocxDocument
# from pdf2image import convert_from_path
# from transformers import pipeline
# from PIL import Image

# # === LangChain Compatibility Patches (for other imports) ===
# if 'langchain.docstore' not in sys.modules:
#     sys.modules['langchain.docstore'] = types.ModuleType('langchain.docstore')
# if 'langchain.docstore.document' not in sys.modules:
#     import langchain_core.documents as lcd
#     sys.modules['langchain.docstore.document'] = types.ModuleType('langchain.docstore.document')
#     sys.modules['langchain.docstore.document'].Document = lcd.Document
# if 'langchain.text_splitter' not in sys.modules:
#     import langchain_text_splitters as lts
#     sys.modules['langchain.text_splitter'] = types.ModuleType('langchain.text_splitter')
#     sys.modules['langchain.text_splitter'].RecursiveCharacterTextSplitter = lts.RecursiveCharacterTextSplitter
# # ===========================================================

# # === Initialize Microsoft TrOCR pipeline ===
# print("🔹 Loading Microsoft TrOCR model... This may take 15–20s on first run.")
# ocr_pipe = pipeline("image-to-text", model="microsoft/trocr-base-stage1")
# print("✅ TrOCR OCR pipeline initialized successfully.")

# def ocr_image(img_path: str) -> str:
#     """Extract text from an image using Microsoft TrOCR pipeline."""
#     try:
#         result = ocr_pipe(Image.open(img_path).convert("RGB"))
#         if isinstance(result, list) and len(result) > 0:
#             return result[0].get("generated_text", "").strip()
#         return ""
#     except Exception as e:
#         print(f"⚠️ OCR failed on {img_path}: {e}")
#         return ""

# def read_pdf(path: str) -> str:
#     """
#     Reads both text-based and scanned PDFs.
#     Uses TrOCR for pages without extractable text.
#     """
#     text = []
#     path = Path(path)
#     if not path.exists():
#         raise FileNotFoundError(f"❌ File not found: {path}")

#     with pdfplumber.open(path) as pdf:
#         for i, page in enumerate(pdf.pages, start=1):
#             extracted = page.extract_text()
#             if extracted and extracted.strip():
#                 text.append(extracted)
#             else:
#                 print(f"🔍 OCR fallback (TrOCR) on page {i} of {path.name}...")
#                 images = convert_from_path(str(path), first_page=i, last_page=i)
#                 for img in images:
#                     temp_path = Path("temp_page.png")
#                     img.save(temp_path)
#                     try:
#                         ocr_text = ocr_image(str(temp_path))
#                         if ocr_text:
#                             text.append(ocr_text)
#                     finally:
#                         temp_path.unlink(missing_ok=True)
#     return "\n".join(text)

# def read_docx(path: str) -> str:
#     """Reads text from Word documents."""
#     doc = DocxDocument(path)
#     return "\n".join(p.text for p in doc.paragraphs)

# def read_text(path: str) -> str:
#     """Reads plain text files."""
#     with open(path, "r", encoding="utf-8") as f:
#         return f.read()

# def clean_text(s: str) -> str:
#     """Cleans up redundant newlines and spaces."""
#     s = s.replace("\r", "\n")
#     s = re.sub(r"\n{3,}", "\n\n", s)
#     s = s.strip()
#     return s

# def list_data_files(data_dir: str):
#     """Lists all PDF, DOCX, and TXT files for ingestion."""
#     p = Path(data_dir)
#     exts = [".pdf", ".docx", ".txt"]
#     return [str(f) for f in p.glob("*") if f.suffix.lower() in exts]


# src/utils.py
# src/utils.py
import os, re, io
from pathlib import Path
import pdfplumber  # pyright: ignore[reportMissingImports]
from docx import Document as DocxDocument  # pyright: ignore[reportMissingImports]
# 🧩 Initialize OCR models (optional - only if available)
trocr_pipe = None
easyocr_reader = None
OCR_AVAILABLE = False
EASYOCR_AVAILABLE = False

try:
    from transformers import pipeline  # pyright: ignore[reportMissingImports]
    from PIL import Image  # pyright: ignore[reportMissingImports]
    import fitz  # PyMuPDF  # pyright: ignore[reportMissingImports]
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import easyocr  # pyright: ignore[reportMissingImports]
    EASYOCR_AVAILABLE = True
except ImportError:
    EASYOCR_AVAILABLE = False

# Initialize OCR engines - OCR is compulsory for processing all PDFs
if OCR_AVAILABLE:
    try:
        print("🧠 Initializing OCR engines (TrOCR + EasyOCR)...")
        try:
            trocr_pipe = pipeline("image-to-text", model="microsoft/trocr-small-printed")
            print("   ✅ TrOCR initialized")
        except Exception as trocr_error:
            print(f"   ⚠️  TrOCR initialization failed: {trocr_error}")
            print("   Will use EasyOCR only")
            trocr_pipe = None
        
        if EASYOCR_AVAILABLE:
            try:
                easyocr_reader = easyocr.Reader(['en'])
                print("   ✅ EasyOCR initialized")
            except Exception as easy_error:
                print(f"   ⚠️  EasyOCR initialization failed: {easy_error}")
                if not trocr_pipe:
                    raise RuntimeError("Both TrOCR and EasyOCR failed to initialize. OCR is required.") from easy_error
        else:
            if not trocr_pipe:
                raise RuntimeError("EasyOCR not available and TrOCR failed. OCR is required.")
        
        print("✅ OCR engines ready.\n")
    except Exception as e:
        print(f"❌ OCR initialization failed: {e}")
        print("   OCR is required for processing all PDFs.")
        print("   Install dependencies: pip install transformers pillow PyMuPDF easyocr")
        raise RuntimeError("OCR initialization failed. OCR is compulsory for processing all PDFs.") from e
else:
    print("❌ OCR dependencies not available. OCR is required for processing all PDFs.")
    print("   Install with: pip install transformers pillow PyMuPDF easyocr")
    raise RuntimeError("OCR dependencies not available. OCR is compulsory for processing all PDFs.")

def read_pdf(path: str) -> str:
    """Extracts text from text-based or scanned PDFs using pdfplumber + TrOCR + EasyOCR fallback."""
    path = Path(path)
    text = []
    if not path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    try:
        with pdfplumber.open(path) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                extracted = page.extract_text()
                if extracted and extracted.strip():
                    text.append(extracted)
                else:
                    # OCR is compulsory - always try OCR when text extraction fails
                    print(f"🔍 OCR fallback on page {i} of {path.name}...")
                    
                    if not OCR_AVAILABLE:
                        raise RuntimeError(
                            f"OCR is required but not available. Page {i} of {path.name} has no extractable text.\n"
                            "Install OCR dependencies: pip install transformers pillow PyMuPDF easyocr"
                        )
                    
                    try:
                        doc = fitz.open(path)
                        page = doc.load_page(i - 1)
                        pix = page.get_pixmap(dpi=200)
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")

                        # Try TrOCR first if available
                        ocr_text = None
                        if trocr_pipe:
                            try:
                                trocr_result = trocr_pipe(img)
                                if trocr_result and len(trocr_result) > 0:
                                    trocr_text = trocr_result[0].get("generated_text", "").strip()
                                    if trocr_text and trocr_text.strip("*") != "":
                                        text.append(trocr_text)
                                        print(f"✅ TrOCR extracted text from page {i}")
                                        ocr_success = True
                                        continue
                            except Exception as trocr_error:
                                print(f"⚠️  TrOCR failed on page {i}: {trocr_error}, trying EasyOCR...")

                        # Fallback to EasyOCR if TrOCR failed or not available
                        if not ocr_success and easyocr_reader:
                            try:
                                easy_results = easyocr_reader.readtext(img_bytes, detail=0)
                                if easy_results:
                                    easy_text = "\n".join(easy_results)
                                    if easy_text.strip():
                                        text.append(easy_text)
                                        print(f"✅ EasyOCR extracted text from page {i}")
                                        ocr_success = True
                                        continue
                            except Exception as easy_error:
                                print(f"⚠️  EasyOCR failed on page {i}: {easy_error}")
                        
                        # If both OCR methods failed, log warning but continue processing
                        # Don't raise error - continue with next page
                        if not ocr_success:
                            print(f"⚠️  Both TrOCR and EasyOCR failed to extract text from page {i} of {path.name}")
                            print(f"   The page might be blank or have unsupported image format. Continuing...")
                            
                    except Exception as ocr_error:
                        error_msg = str(ocr_error)
                        if "OCR is required" in error_msg or "Both TrOCR and EasyOCR failed" in error_msg:
                            raise  # Re-raise critical errors
                        print(f"⚠️  OCR error on page {i}: {error_msg}")
                        # Try to continue with next page, but log the error
                        print(f"   Skipping page {i} due to OCR failure")

    except Exception as e:
        print(f"⚠️ PDF extraction failed ({e}).")
        if OCR_AVAILABLE:
            print("   Attempting full OCR on all pages...")
            try:
                doc = fitz.open(path)
                for i, page in enumerate(doc, start=1):
                    try:
                        pix = page.get_pixmap(dpi=200)
                        img_bytes = pix.tobytes("png")
                        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")
                        
                        # Try TrOCR first
                        ocr_text = None
                        if trocr_pipe:
                            try:
                                trocr_result = trocr_pipe(img)
                                if trocr_result and len(trocr_result) > 0:
                                    ocr_text = trocr_result[0].get("generated_text", "").strip()
                                    if ocr_text and ocr_text.strip("*") != "":
                                        text.append(ocr_text)
                                        print(f"✅ TrOCR extracted text from page {i}")
                                        continue
                            except:
                                pass
                        
                        # Fallback to EasyOCR
                        if not ocr_text and easyocr_reader:
                            try:
                                easy_results = easyocr_reader.readtext(img_bytes, detail=0)
                                if easy_results:
                                    easy_text = "\n".join(easy_results)
                                    if easy_text.strip():
                                        text.append(easy_text)
                                        print(f"✅ EasyOCR extracted text from page {i}")
                            except Exception as easy_error:
                                print(f"⚠️  EasyOCR failed on page {i}: {easy_error}")
                    except Exception as page_error:
                        print(f"⚠️  Failed to process page {i}: {page_error}")
            except Exception as ocr_error:
                print(f"⚠️  Full OCR also failed: {ocr_error}")
                raise RuntimeError(f"Failed to extract text from {path.name} using both text extraction and OCR") from ocr_error
        else:
            raise RuntimeError(
                f"Failed to extract text from {path.name} and OCR is not available.\n"
                "Install OCR dependencies: pip install transformers pillow PyMuPDF easyocr"
            ) from e

    full_text = "\n".join(text)
    if full_text.strip():
        print(f"✅ Extracted {len(full_text)} characters from {path.name}")
    else:
        print(f"⚠️  No text extracted from {path.name}")
    return full_text


def read_docx(path: str) -> str:
    if DocxDocument is None:
        raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
    doc = DocxDocument(path)
    return "\n".join(p.text for p in doc.paragraphs)


def read_text(path: str) -> str:
    with open(path, "r", encoding="utf-8") as f:
        return f.read()


def clean_text(s: str) -> str:
    s = s.replace("\r", "\n")
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def list_data_files(data_dir: str):
    """
    Returns a list of valid document file paths (.pdf, .docx, .txt)
    inside the given directory.
    """
    p = Path(data_dir)
    exts = [".pdf", ".docx", ".txt"]
    files = [str(f) for f in p.glob("*") if f.suffix.lower() in exts]
    
    if not files:
        print(f"⚠️ No data files found in {data_dir}")
    else:
        print(f"📂 Found {len(files)} data files in {data_dir}")
    
    return files

